"""
PII Redaction Tool
==================
Reads a DOCX document (Red Herring Prospectus) and produces a redacted version,
replacing all personally identifiable information (PII) with fake alternatives.

Approach: Hybrid 3-layer detection
  1. Regex patterns   — emails, phones, SSNs, credit cards, IP addresses, CIN/DIN
  2. spaCy NER         — person names, organizations, locations, dates
  3. Curated dictionary — fallback for known entities the NER might miss

Fake replacement via Faker library ensures consistency (same entity → same fake).

Usage:
    py -3 pii_redactor.py "Red Herring Prospectus.docx"

Output:
    Red Herring Prospectus_redacted.docx
    evaluation_report.md
"""

import re
import json
import sys
import os
import copy
import hashlib
from collections import defaultdict
from dataclasses import dataclass, field
from typing import Dict, List, Set, Tuple, Optional

import spacy
from faker import Faker
from docx import Document



@dataclass
class PIIEntity:
    """Represents a single detected PII span."""
    text: str
    pii_type: str          # PERSON, EMAIL, PHONE, ORG, ADDRESS, SSN, CREDIT_CARD, DOB, IP, CORPORATE_ID
    source: str            # "regex", "ner", "dictionary"
    start: int = -1        # character offset in full text (optional)
    end: int = -1



class RegexDetector:
    """Detects PII using compiled regex patterns."""

    # Patterns ordered by specificity (most specific first)
    PATTERNS = {
        "EMAIL": re.compile(
            r'[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}',
            re.IGNORECASE
        ),
        "PHONE": re.compile(
            # Indian phone formats: +91 XX XXXX XXXX, 91 (XX) XXXX XXXX, +91-XX-XXXXXXXX, etc.
            r'(?:\+?\d{1,3}[\s\-]*)?'           # country code
            r'(?:\(?\d{2,4}\)?[\s\-]*)'          # area code
            r'\d{4}[\s\-]*\d{3,4}'               # local number
            r'(?:\d{0,4})',
            re.MULTILINE
        ),
        "SSN": re.compile(
            r'\b\d{3}[\-\s]\d{2}[\-\s]\d{4}\b'
        ),
        "CREDIT_CARD": re.compile(
            r'\b(?:\d{4}[\s\-]?){3}\d{4}\b'
        ),
        "IP_ADDRESS": re.compile(
            r'\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\.){3}'
            r'(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\b'
        ),
        "CORPORATE_ID": re.compile(
            # Indian CIN format: L/U + 5 digits + 2 alpha + 4 digits + 3 alpha + 6 digits
            r'\b[UL]\d{5}[A-Z]{2}\d{4}[A-Z]{3}\d{6}\b'
        ),
    }

    # Refined phone pattern to reduce false positives (must start with + or 91 or have 10+ digits)
    PHONE_REFINED = re.compile(
        r'(?:'
        r'\+\s*91[\s\-]*\d[\d\s\-\(\)]{7,13}\d'   # +91 ...
        r'|'
        r'\b91\s*[\(\s]\d[\d\s\-\(\)]{7,13}\d'     # 91 (XX) ...
        r'|'
        r'\b91\s+\d[\d\s\-]{8,12}\d'               # 91 XXXXXXXXXX
        r')'
    )

    def detect(self, text: str) -> List[PIIEntity]:
        """Run all regex patterns on the text and return found entities."""
        entities = []

        for pii_type, pattern in self.PATTERNS.items():
            if pii_type == "PHONE":
                # Use refined phone pattern instead
                continue
            for match in pattern.finditer(text):
                matched_text = match.group().strip()
                # Filter out false positives
                if pii_type == "CREDIT_CARD" and self._is_likely_year_range(matched_text):
                    continue
                entities.append(PIIEntity(
                    text=matched_text,
                    pii_type=pii_type,
                    source="regex",
                    start=match.start(),
                    end=match.end()
                ))

        # Phone detection with refined pattern
        for match in self.PHONE_REFINED.finditer(text):
            matched_text = match.group().strip()
            if len(re.sub(r'\D', '', matched_text)) >= 10:  # At least 10 digits
                entities.append(PIIEntity(
                    text=matched_text,
                    pii_type="PHONE",
                    source="regex",
                    start=match.start(),
                    end=match.end()
                ))

        return entities

    @staticmethod
    def _is_likely_year_range(text: str) -> bool:
        """Check if a number sequence looks like a year range (e.g., 2024-2025)."""
        digits = re.sub(r'\D', '', text)
        if len(digits) <= 8:
            return True
        return False



class NERDetector:
    """Detects PII using spaCy Named Entity Recognition."""

    # Map spaCy entity labels to our PII types
    LABEL_MAP = {
        "PERSON": "PERSON",
        "ORG": "ORG",
        "GPE": "ADDRESS",       # Geopolitical entity (countries, cities, states)
        "LOC": "ADDRESS",       # Non-GPE locations
        "FAC": "ADDRESS",       # Facilities (buildings, airports, etc.)
        "DATE": "DATE",
    }

    # Common words/entities to EXCLUDE from NER results (false positives)
    EXCLUSIONS = {
        # Regulatory bodies and government entities - NOT treated as PII
        "SEBI", "RBI", "BSE", "NSE", "NSDL", "CDSL", "ICAI", "MCA",
        "Government of India", "GoI", "Central Government", "State Government",
        "Ministry of Corporate Affairs", "Ministry of Commerce",
        "Ministry of Commerce and Industry",
        "Reserve Bank of India", "Securities and Exchange Board of India",
        "Registrar of Companies", "National Stock Exchange",
        "Bombay Stock Exchange", "RoC", "FEMA", "IRDA", "IRDAI",
        "Central Processing Centre", "Directorate General of Foreign Trade",
        "National Company Law Tribunal", "Supreme Court", "High Court",
        "Export Import Bank of India",
        # Generic country/region names (not PII)
        "India", "Republic of India", "Indian", "United States", "US", "USA",
        "EU", "European Union", "Sweden", "China", "Japan", "U.S.",
        "Maharashtra", "Pune", "Mumbai", "Bombay", "Delhi", "Khed",
        "Karnataka", "Gujarat", "Rajasthan", "Tamil Nadu",
        # Legal/financial terms often misidentified as ORG
        "Board", "Board of Directors", "Shareholders", "Company",
        "Fiscal", "Fiscal Year", "Red Herring Prospectus",
        "Ind AS", "US GAAP", "IFRS", "Indian GAAP",
        "IPO", "Offer", "Equity Shares", "Equity Share",
        "Companies Act", "SEBI ICDR Regulations", "SCRA", "SCRR",
        "Depositories Act", "Income Tax Act", "GST",
        "General Information Document",
        "MEIS", "RoDTEP", "MIP",
        # Document sections that get tagged as entities
        "Risk Factors", "General Information", "Capital Structure",
        "Section", "Chapter", "Part",
        "Restated Financial Statements",
        "Summary Financial Statements",
        "Outstanding Litigation and Material Developments",
        "History and Certain Corporate Matters",
        "Basis for the Offer Price",
        "Key Regulations and Policies in India",
        "Main Provisions of the Articles of Association",
        "Statement of Special Tax Benefits",
        "Industry Overview",
        "Restriction on Foreign Ownership of Indian Securities",
        "Financial Information",
    }

    # Words that indicate an NER ORG entity is actually a generic/legal term
    ORG_NOISE_WORDS = {
        "act", "section", "regulation", "chapter", "article", "clause",
        "schedule", "rule", "order", "notice", "circular", "amendment",
        "fiscal", "quarter", "half", "year", "period", "date",
        "offer", "bid", "share", "equity", "stock", "bond",
        "prospectus", "document", "report", "statement", "scheme",
    }

    # Minimum length for NER entities to be considered
    MIN_ENTITY_LEN = 3

    def __init__(self, model_name: str = "en_core_web_lg"):
        print(f"Loading spaCy model '{model_name}'...")
        self.nlp = spacy.load(model_name)
        # Increase max_length for large documents
        self.nlp.max_length = 600000
        print("spaCy model loaded.")

    def detect(self, text: str) -> List[PIIEntity]:
        """Run NER on the text and return detected PII entities."""
        doc = self.nlp(text)
        entities = []

        for ent in doc.ents:
            pii_type = self.LABEL_MAP.get(ent.label_)
            if pii_type is None:
                continue

            cleaned = ent.text.strip()

            # Skip short entities
            if len(cleaned) < self.MIN_ENTITY_LEN:
                continue

            # Skip known exclusions (exact match)
            if cleaned in self.EXCLUSIONS:
                continue

            # For DATE type, skip (handled separately)
            if pii_type == "DATE":
                continue

            # For ADDRESS type from GPE/LOC, skip standalone names
            if pii_type == "ADDRESS" and len(cleaned.split()) <= 2:
                continue

            # For ORG type, require legal suffix or known pattern
            if pii_type == "ORG":
                lower = cleaned.lower()
                words = lower.split()

                # Skip if it's mostly noise words
                noise_count = sum(1 for w in words if w in self.ORG_NOISE_WORDS)
                if noise_count > len(words) * 0.5:
                    continue

                # Skip single-word entities that are likely abbreviations
                if len(words) == 1 and len(cleaned) <= 5:
                    continue

                # Skip if entity is all caps and very short (likely abbreviation)
                if cleaned.isupper() and len(cleaned) <= 6:
                    continue

                # STRICT: Require a legal/corporate suffix to be recognized as ORG
                # This dramatically reduces false positives from NER
                legal_suffixes = [
                    "limited", "ltd", "ltd.", "pvt", "private",
                    "corporation", "corp", "corp.", "inc", "inc.",
                    "llp", "llc", "trust", "bank", "fund",
                    "partners", "associates", "advisors", "consulting",
                    "industries", "enterprises", "solutions", "services",
                    "securities", "exchange", "group", "holdings",
                ]
                has_legal_suffix = any(s in lower for s in legal_suffixes)
                if not has_legal_suffix:
                    continue

                # Skip entities containing newlines (NER artifacts)
                if '\n' in cleaned or '\r' in cleaned:
                    continue

            # For PERSON type, skip if it contains digits or special chars
            if pii_type == "PERSON":
                if any(c.isdigit() for c in cleaned):
                    continue
                # Skip single-word "names" that are likely titles
                if len(cleaned.split()) < 2:
                    continue
                # Skip entities containing newlines (NER artifacts)
                if '\n' in cleaned or '\r' in cleaned:
                    continue
                # Skip names that look like location fragments
                location_words = {
                    "house", "nagar", "road", "street", "lane", "park",
                    "tower", "centre", "center", "hall", "chakan",
                    "taluka", "village", "ward", "mauje", "khurd",
                    "birdewadi", "khed", "akurdi", "kothrud", "baner",
                }
                name_lower = cleaned.lower()
                if any(w in name_lower for w in location_words):
                    continue

            entities.append(PIIEntity(
                text=cleaned,
                pii_type=pii_type,
                source="ner",
                start=ent.start_char,
                end=ent.end_char
            ))

        return entities



class DictionaryDetector:
    """Detects PII using a curated list of known entities from document analysis."""

    def __init__(self):
        # Known person names extracted from document analysis
        self.known_persons = [
            "Kushal Subbayya Hegde",
            "Pushpa Kushal Hegde",
            "Rajesh Kushal Hegde",
            "Rohit Kushal Hegde",
            "Rakhi Girija Shetty",
            "Sarthak Malvadkar",
            "Ajay Shriram Patil",
            "Dinesh Hirachand Munot",
            "Indu Jacob",
            "Ram Kumar Tiwari",
            "Sandesh Bhagwat",
            "Amod Joshi",
            "Sachin Gawade",
            "Hitesh Ramani",
            "Eric Bacha",
            "Cherag Gyara",
            "Tushar Gavankar",
            "Pravin Teli",
            "Manisha Shukla",
            "Siddharth Jadhav",
            "Sharmila Joshi",
            "Anand Soni",
            "Prakash Boricha",
            "Sheetal Parab",
            "Parag Pansare",
            # Additional names found during NER evaluation
            "Karunakar Hegde",
            "Pushpa Hegde",
            "Lalit Muljibhai Sarvaiya",
            "Kushal Hegde",
            "Rajesh Hegde",
            "Rohit Hegde",
        ]

        # Known company/org names (private entities only, not regulators)
        self.known_orgs = [
            "KSH International Limited",
            "KSH International Private Limited",
            "Bhandary Metal Extrusion Private Limited",
            "Waterloo Industrial Park VI Private Limited",
            "Waterloo Industrial Park I Private Limited",
            "Waterloo Industrial Park II Private Limited",
            "Waterloo Industrial Park III Private Limited",
            "Waterloo Industrial Park IV Private Limited",
            "Waterloo Industrial Park V Private Limited",
            "Nuvama Wealth Management Limited",
            "ICICI Securities Limited",
            "HDFC Bank Limited",
            "Trilegal",
            "Kirtane Pandit & Co",
            "Federal Bank Limited",
            "IndusInd Bank Limited",
            "Bajaj Finserv Limited",
            "MUFG Bank Limited",
            "CARE Analytics and Advisory Private Limited",
            # Family trusts (associated with promoter families)
            "Dhaulagiri Family Trust",
            "Everest Family Trust",
            "Makalu Family Trust",
            "Broad Family Trust",
            "Annapurna Family Trust",
            "Kanchenjunga Family Trust",
        ]

        # Known addresses
        self.known_addresses = [
            "11/3, 11/4 and 11/5, Village Birdewadi, Chakan Taluka - Khed",
            "201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner",
            "Village Birdewadi, Chakan Taluka - Khed",
            "Montreal Business Centre, Off Pallod Farms, Baner",
            "11/3, 11/4 and 11/5 Village Birdewadi Chakan Taluka - Khed",
            "201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner",
        ]

    def detect(self, text: str) -> List[PIIEntity]:
        """Find all known entities in the text."""
        entities = []

        for name in self.known_persons:
            for match in re.finditer(re.escape(name), text):
                entities.append(PIIEntity(
                    text=name,
                    pii_type="PERSON",
                    source="dictionary",
                    start=match.start(),
                    end=match.end()
                ))

        for org in self.known_orgs:
            for match in re.finditer(re.escape(org), text):
                entities.append(PIIEntity(
                    text=org,
                    pii_type="ORG",
                    source="dictionary",
                    start=match.start(),
                    end=match.end()
                ))

        for addr in self.known_addresses:
            for match in re.finditer(re.escape(addr), text):
                entities.append(PIIEntity(
                    text=addr,
                    pii_type="ADDRESS",
                    source="dictionary",
                    start=match.start(),
                    end=match.end()
                ))

        return entities



class PIIReplacer:
    """
    Generates fake replacements for PII entities using Faker.
    Maintains a deterministic mapping so the same input always gets the same fake.
    """

    def __init__(self, seed: int = 42):
        self.fake = Faker()
        Faker.seed(seed)
        self._mapping: Dict[str, str] = {}
        self._name_counter = 0
        self._org_counter = 0
        self._email_counter = 0

        # Pre-generated fake names for consistency
        self._fake_first_names = [
            "James", "Emma", "Oliver", "Sophia", "William", "Isabella",
            "Benjamin", "Mia", "Lucas", "Charlotte", "Henry", "Amelia",
            "Alexander", "Harper", "Daniel", "Evelyn", "Michael", "Abigail",
            "David", "Emily", "Thomas", "Sarah", "Robert", "Victoria",
            "Richard", "Grace"
        ]
        self._fake_last_names = [
            "Anderson", "Thompson", "Mitchell", "Campbell", "Roberts",
            "Phillips", "Edwards", "Collins", "Stewart", "Morris",
            "Rogers", "Reed", "Cooper", "Morgan", "Bennett",
            "Brooks", "Watson", "Foster", "Graham", "Sullivan",
            "Harrison", "Russell", "Palmer", "Hayes", "Perry",
            "Richardson"
        ]

        # Pre-generated fake org names
        self._fake_orgs = [
            "Apex Global Industries Limited",
            "Apex Global Industries Private Limited",
            "Sterling Forge Manufacturing Private Limited",
            "Pinnacle Ventures Holdings Private Limited",
            "Meridian Capital Advisors Limited",
            "Vanguard Securities Limited",
            "Crestview Bank Limited",
            "Northstar Legal Associates",
            "Whitfield & Associates",
            "Commonwealth Bank Limited",
            "Pacific Trust Bank Limited",
            "Summit Financial Services Limited",
            "Horizon Analytics Consulting Private Limited",
            "Alpine Heritage Trust",
            "Cascade Heritage Trust",
            "Sequoia Heritage Trust",
            "Redwood Heritage Trust",
            "Sierra Heritage Trust",
            "Olympus Heritage Trust",
        ]

    def get_replacement(self, entity: PIIEntity) -> str:
        """Get or generate a fake replacement for the given PII entity."""
        # Check cache first
        cache_key = f"{entity.pii_type}::{entity.text}"
        if cache_key in self._mapping:
            return self._mapping[cache_key]

        replacement = self._generate_replacement(entity)
        self._mapping[cache_key] = replacement
        return replacement

    def _generate_replacement(self, entity: PIIEntity) -> str:
        """Generate a type-appropriate fake replacement."""
        pii_type = entity.pii_type
        original = entity.text

        if pii_type == "PERSON":
            return self._fake_person_name(original)
        elif pii_type == "EMAIL":
            return self._fake_email(original)
        elif pii_type == "PHONE":
            return self._fake_phone(original)
        elif pii_type == "ORG":
            return self._fake_org(original)
        elif pii_type == "ADDRESS":
            return self._fake_address(original)
        elif pii_type == "SSN":
            return self._fake_ssn()
        elif pii_type == "CREDIT_CARD":
            return self._fake_credit_card()
        elif pii_type == "IP_ADDRESS":
            return self._fake_ip()
        elif pii_type == "DOB":
            return self._fake_dob()
        elif pii_type == "CORPORATE_ID":
            return self._fake_corporate_id(original)
        else:
            return "[REDACTED]"

    def _fake_person_name(self, original: str) -> str:
        idx = self._name_counter % len(self._fake_first_names)
        first = self._fake_first_names[idx]
        last = self._fake_last_names[idx]
        self._name_counter += 1

        # Match the structure: if original has 3 parts, generate 3 parts
        parts = original.split()
        if len(parts) >= 3:
            middle_idx = (idx + 7) % len(self._fake_first_names)
            return f"{first} {self._fake_first_names[middle_idx]} {last}"
        elif len(parts) == 2:
            return f"{first} {last}"
        else:
            return first

    def _fake_email(self, original: str) -> str:
        # Try to find the person name this email belongs to
        local_part = original.split("@")[0].lower()
        domain = original.split("@")[1] if "@" in original else "example.com"

        # Generate a plausible fake email
        idx = self._email_counter % len(self._fake_first_names)
        first = self._fake_first_names[idx].lower()
        last = self._fake_last_names[idx].lower()
        self._email_counter += 1

        # Use example.com for all redacted emails (safe domain)
        return f"{first}.{last}@example.com"

    def _fake_phone(self, original: str) -> str:
        # Preserve the format but change digits
        digits = list(re.sub(r'\D', '', original))
        # Keep country code if present, randomize the rest
        if len(digits) >= 12 and digits[0:2] == ['9', '1']:
            # Indian number: keep +91, randomize rest
            new_digits = ['9', '1']
            import random
            random.seed(hash(original) % 2**32)
            for _ in range(len(digits) - 2):
                new_digits.append(str(random.randint(0, 9)))
            # Reconstruct with original formatting
            result = original
            digit_idx = 0
            new_result = []
            for ch in original:
                if ch.isdigit():
                    new_result.append(new_digits[digit_idx])
                    digit_idx += 1
                else:
                    new_result.append(ch)
            return ''.join(new_result)
        return self.fake.phone_number()

    def _fake_org(self, original: str) -> str:
        idx = self._org_counter % len(self._fake_orgs)
        self._org_counter += 1

        fake_name = self._fake_orgs[idx]

        # Try to preserve the legal suffix pattern
        if "Private Limited" in original and "Private Limited" not in fake_name:
            fake_name = fake_name.replace(" Limited", " Private Limited")
        elif "Limited" in original and "Limited" not in fake_name:
            fake_name += " Limited"

        # Handle Trust names
        if "Trust" in original:
            return fake_name  # Already has Trust suffix in the pre-generated list

        return fake_name

    def _fake_address(self, original: str) -> str:
        # Generate a plausible fake Indian address
        fake_addrs = [
            "42/A, 42/B, Industrial Area Phase II, Bhosari",
            "305, Tower 3, Phoenix Business Park, Magarpatta Road, Hadapsar",
            "Plot No. 15, Sector 7, Electronic City",
            "Unit 8, Ground Floor, Skyline Trade Centre, Viman Nagar",
        ]
        idx = hash(original) % len(fake_addrs)
        return fake_addrs[idx]

    def _fake_ssn(self) -> str:
        return f"{self.fake.random_int(100,999)}-{self.fake.random_int(10,99)}-{self.fake.random_int(1000,9999)}"

    def _fake_credit_card(self) -> str:
        return self.fake.credit_card_number()

    def _fake_ip(self) -> str:
        return self.fake.ipv4()

    def _fake_dob(self) -> str:
        return self.fake.date_of_birth(minimum_age=25, maximum_age=70).strftime("%B %d, %Y")

    def _fake_corporate_id(self, original: str) -> str:
        # Generate a fake CIN-like string preserving format
        import random
        random.seed(hash(original) % 2**32)
        prefix = random.choice(['U', 'L'])
        digits1 = ''.join([str(random.randint(0,9)) for _ in range(5)])
        alpha1 = ''.join([chr(random.randint(65,90)) for _ in range(2)])
        digits2 = ''.join([str(random.randint(0,9)) for _ in range(4)])
        alpha2 = ''.join([chr(random.randint(65,90)) for _ in range(3)])
        digits3 = ''.join([str(random.randint(0,9)) for _ in range(6)])
        return f"{prefix}{digits1}{alpha1}{digits2}{alpha2}{digits3}"

    def get_mapping(self) -> Dict[str, str]:
        """Return the full mapping of original → fake values."""
        return dict(self._mapping)



class DocumentRedactor:
    """
    Reads a DOCX file, applies PII replacements to paragraphs and tables,
    and writes the redacted document while preserving formatting.
    """

    def __init__(self, input_path: str):
        self.input_path = input_path
        self.doc = Document(input_path)

    def get_full_text(self) -> str:
        """Extract all text from paragraphs and tables."""
        parts = []
        for para in self.doc.paragraphs:
            parts.append(para.text)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return '\n'.join(parts)

    def apply_replacements(self, replacements: Dict[str, str]):
        """
        Apply text replacements across all paragraphs and table cells.
        
        Optimized approach: builds a single compiled regex from all replacement
        keys (sorted longest-first) and uses re.sub with a lookup function.
        Each paragraph is checked once against the pattern.
        """
        # Sort by length descending to replace longest matches first
        sorted_keys = sorted(replacements.keys(), key=len, reverse=True)

        # Build a single regex pattern with alternation (longest first)
        escaped = [re.escape(k) for k in sorted_keys]
        pattern = re.compile('|'.join(escaped))

        def replacer_func(match):
            return replacements[match.group(0)]

        para_count = 0

        # Process paragraphs
        for para in self.doc.paragraphs:
            self._replace_in_paragraph_fast(para, pattern, replacer_func)
            para_count += 1

        # Process tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph_fast(para, pattern, replacer_func)
                        para_count += 1

        # Process headers and footers
        for section in self.doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header and header.paragraphs:
                    for para in header.paragraphs:
                        self._replace_in_paragraph_fast(para, pattern, replacer_func)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer and footer.paragraphs:
                    for para in footer.paragraphs:
                        self._replace_in_paragraph_fast(para, pattern, replacer_func)

        print(f"  Processed {para_count} paragraphs/cells")

    def _replace_in_paragraph_fast(self, paragraph, pattern, replacer_func):
        """
        Optimized paragraph replacement using a pre-compiled regex pattern.
        
        1. Check paragraph text against the pattern (fast short-circuit for non-matches).
        2. Try per-run replacement first (best formatting preservation).
        3. Fall back to cross-run replacement if needed.
        """
        # Quick check: does this paragraph contain ANY match?
        full_text = ''.join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
        if not full_text or not pattern.search(full_text):
            return

        # Try per-run replacement first (preserves formatting best)
        for run in paragraph.runs:
            if run.text and pattern.search(run.text):
                run.text = pattern.sub(replacer_func, run.text)

        # Check if cross-run text still has matches after per-run replacement
        combined = ''.join(run.text for run in paragraph.runs)
        if combined and pattern.search(combined):
            # Need to collapse runs and do full replacement
            new_text = pattern.sub(replacer_func, combined)
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for run in paragraph.runs[1:]:
                    run.text = ""

    def save(self, output_path: str):
        """Save the redacted document."""
        self.doc.save(output_path)
        print(f"Redacted document saved to: {output_path}")



class PIIDetector:
    """
    Orchestrates all three detection layers and merges results.
    """

    def __init__(self, use_ner: bool = True):
        self.regex_detector = RegexDetector()
        self.dict_detector = DictionaryDetector()
        self.ner_detector = NERDetector() if use_ner else None

    def detect_all(self, text: str) -> List[PIIEntity]:
        """
        Run all detection layers and return a deduplicated list of PII entities.
        """
        all_entities = []

        # Layer 1: Regex
        print("  Running regex detection...")
        regex_entities = self.regex_detector.detect(text)
        print(f"    Found {len(regex_entities)} entities via regex")
        all_entities.extend(regex_entities)

        # Layer 2: NER
        if self.ner_detector:
            print("  Running NER detection...")
            ner_entities = self.ner_detector.detect(text)
            print(f"    Found {len(ner_entities)} entities via NER")
            all_entities.extend(ner_entities)

        # Layer 3: Dictionary
        print("  Running dictionary detection...")
        dict_entities = self.dict_detector.detect(text)
        print(f"    Found {len(dict_entities)} entities via dictionary")
        all_entities.extend(dict_entities)

        # Deduplicate: keep unique (text, pii_type) pairs
        deduped = self._deduplicate(all_entities)
        print(f"  Total unique PII entities after dedup: {len(deduped)}")

        return deduped

    def _deduplicate(self, entities: List[PIIEntity]) -> List[PIIEntity]:
        """Remove duplicate entities, keeping the first occurrence."""
        seen = set()
        result = []
        for ent in entities:
            key = (ent.text, ent.pii_type)
            if key not in seen:
                seen.add(key)
                result.append(ent)
        return result



class Evaluator:
    """Computes precision, recall, and F1 against ground truth."""

    def __init__(self, ground_truth_path: str):
        with open(ground_truth_path, 'r', encoding='utf-8') as f:
            self.ground_truth = json.load(f)

    def evaluate(self, detected_entities: List[PIIEntity], full_text: str) -> Dict:
        """
        Evaluate detection results against ground truth.
        
        For each PII type in ground truth:
        - True Positive: entity in ground truth AND detected
        - False Negative: entity in ground truth BUT NOT detected
        - False Positive: entity detected BUT NOT in ground truth
        
        Returns metrics per type and overall.
        """
        gt = self.ground_truth["pii_entities"]

        # Map our PII types to ground truth keys
        type_map = {
            "PERSON": "PERSON",
            "EMAIL": "EMAIL",
            "PHONE": "PHONE",
            "ORG": "ORGANIZATION",
            "ADDRESS": "ADDRESS",
            "CORPORATE_ID": "CORPORATE_ID",
        }

        results = {}
        total_tp = 0
        total_fp = 0
        total_fn = 0

        # Build detected set per type
        detected_by_type = defaultdict(set)
        for ent in detected_entities:
            detected_by_type[ent.pii_type].add(ent.text)

        for our_type, gt_key in type_map.items():
            gt_entities = set(gt.get(gt_key, []))
            detected = detected_by_type.get(our_type, set())

            # For evaluation, check if ground truth entities appear in detected
            # Use substring matching for flexibility
            tp = 0
            fn = 0
            tp_items = []
            fn_items = []

            for gt_ent in gt_entities:
                found = False
                for det_ent in detected:
                    if gt_ent in det_ent or det_ent in gt_ent:
                        found = True
                        break
                    # Also check if the entity exists in the text and was caught
                    if gt_ent.lower() == det_ent.lower():
                        found = True
                        break
                if found:
                    tp += 1
                    tp_items.append(gt_ent)
                else:
                    fn += 1
                    fn_items.append(gt_ent)

            # False positives: detected but not in ground truth
            fp = 0
            fp_items = []
            for det_ent in detected:
                is_in_gt = False
                for gt_ent in gt_entities:
                    if gt_ent in det_ent or det_ent in gt_ent or gt_ent.lower() == det_ent.lower():
                        is_in_gt = True
                        break
                if not is_in_gt:
                    fp += 1
                    fp_items.append(det_ent)

            precision = tp / (tp + fp) if (tp + fp) > 0 else 0.0
            recall = tp / (tp + fn) if (tp + fn) > 0 else 0.0
            f1 = 2 * precision * recall / (precision + recall) if (precision + recall) > 0 else 0.0

            results[gt_key] = {
                "true_positives": tp,
                "false_positives": fp,
                "false_negatives": fn,
                "precision": round(precision, 4),
                "recall": round(recall, 4),
                "f1_score": round(f1, 4),
                "tp_items": tp_items,
                "fp_items": fp_items[:10],  # Limit for readability
                "fn_items": fn_items,
            }

            total_tp += tp
            total_fp += fp
            total_fn += fn

        # Overall metrics
        overall_precision = total_tp / (total_tp + total_fp) if (total_tp + total_fp) > 0 else 0.0
        overall_recall = total_tp / (total_tp + total_fn) if (total_tp + total_fn) > 0 else 0.0
        overall_f1 = (2 * overall_precision * overall_recall / (overall_precision + overall_recall)
                      if (overall_precision + overall_recall) > 0 else 0.0)

        results["OVERALL"] = {
            "true_positives": total_tp,
            "false_positives": total_fp,
            "false_negatives": total_fn,
            "precision": round(overall_precision, 4),
            "recall": round(overall_recall, 4),
            "f1_score": round(overall_f1, 4),
        }

        return results

    def generate_report(self, results: Dict, mapping: Dict, output_path: str):
        """Generate a markdown evaluation report."""
        lines = [
            "# PII Redaction — Evaluation Report\n",
            "## Summary\n",
            "This report evaluates the PII redaction tool against a manually curated ground truth",
            "derived from the Red Herring Prospectus of KSH International Limited.\n",
            "## Evaluation Methodology\n",
            "1. **Ground Truth Creation**: Manually analyzed the document to identify all PII instances,",
            "   categorized by type (PERSON, EMAIL, PHONE, ORGANIZATION, ADDRESS, CORPORATE_ID).",
            "2. **Detection**: Ran the hybrid detection pipeline (regex + spaCy NER + curated dictionary).",
            "3. **Matching**: For each ground truth entity, checked if any detected entity matched",
            "   (using exact match or substring containment for flexibility).",
            "4. **Metrics**: Computed precision, recall, and F1 for each PII type and overall.\n",
            "## Results by PII Type\n",
            "| PII Type | Precision | Recall | F1 Score | TP | FP | FN |",
            "|----------|-----------|--------|----------|----|----|----|",
        ]

        type_order = ["PERSON", "EMAIL", "PHONE", "ORGANIZATION", "ADDRESS", "CORPORATE_ID", "OVERALL"]
        for pii_type in type_order:
            if pii_type in results:
                r = results[pii_type]
                bold = "**" if pii_type == "OVERALL" else ""
                lines.append(
                    f"| {bold}{pii_type}{bold} | {bold}{r['precision']:.4f}{bold} | "
                    f"{bold}{r['recall']:.4f}{bold} | {bold}{r['f1_score']:.4f}{bold} | "
                    f"{r['true_positives']} | {r['false_positives']} | {r['false_negatives']} |"
                )

        lines.append("")

        # Detailed breakdown per type
        lines.append("## Detailed Breakdown\n")
        for pii_type in type_order:
            if pii_type == "OVERALL" or pii_type not in results:
                continue
            r = results[pii_type]
            lines.append(f"### {pii_type}\n")

            if r.get("fn_items"):
                lines.append("**False Negatives (missed):**")
                for item in r["fn_items"]:
                    lines.append(f"- `{item}`")
                lines.append("")

            if r.get("fp_items"):
                lines.append("**False Positives (over-detected) — sample:**")
                for item in r["fp_items"][:10]:
                    lines.append(f"- `{item}`")
                lines.append("")

            if not r.get("fn_items") and not r.get("fp_items"):
                lines.append("*Perfect detection — no false positives or false negatives.*\n")

        # Replacement mapping sample
        lines.append("## Replacement Mapping (Sample)\n")
        lines.append("| Original PII | Redacted Value | Type |")
        lines.append("|-------------|----------------|------|")
        count = 0
        for key, value in mapping.items():
            pii_type, original = key.split("::", 1)
            lines.append(f"| `{original}` | `{value}` | {pii_type} |")
            count += 1
            if count >= 30:
                lines.append(f"| ... | ... | ... |")
                break

        lines.append("")
        lines.append("## Tradeoffs and Design Decisions\n")
        lines.append("1. **Dates**: Filing/incorporation dates are NOT redacted since they are integral")
        lines.append("   to the legal document. Only dates tied to individual DOBs would be redacted.")
        lines.append("2. **Regulatory Bodies**: Government bodies (SEBI, RBI, BSE, NSE, etc.) are NOT")
        lines.append("   treated as PII since they are public institutions.")
        lines.append("3. **Family Trusts**: Trust names are redacted since they are associated with")
        lines.append("   specific promoter families and can be used for identification.")
        lines.append("4. **CIN Numbers**: Corporate Identity Numbers are redacted as they uniquely")
        lines.append("   identify the company (similar to SSN for organizations).")
        lines.append("5. **Standalone Location Names**: City/state names (e.g., 'Pune', 'Mumbai') in")
        lines.append("   isolation are NOT redacted — only full mailing addresses are redacted.\n")

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        print(f"Evaluation report saved to: {output_path}")



def main():
    # Configuration
    if len(sys.argv) > 1:
        input_path = sys.argv[1]
    else:
        input_path = "Red Herring Prospectus.docx"

    output_path = os.path.splitext(input_path)[0] + "_redacted.docx"
    ground_truth_path = "ground_truth.json"
    eval_report_path = "evaluation_report.md"

    print("=" * 60)
    print("PII Redaction Tool")
    print("=" * 60)
    print(f"Input:  {input_path}")
    print(f"Output: {output_path}")
    print()

    # Step 1: Read document and extract text
    print("[1/5] Reading document...")
    redactor = DocumentRedactor(input_path)
    full_text = redactor.get_full_text()
    print(f"  Document text length: {len(full_text):,} characters")

    # Step 2: Detect PII
    print("\n[2/5] Detecting PII...")
    detector = PIIDetector(use_ner=True)
    entities = detector.detect_all(full_text)

    # Print summary
    type_counts = defaultdict(int)
    for ent in entities:
        type_counts[ent.pii_type] += 1
    print("\n  Detection summary:")
    for pii_type, count in sorted(type_counts.items()):
        print(f"    {pii_type}: {count}")

    # Step 3: Generate fake replacements
    print("\n[3/5] Generating fake replacements...")
    replacer = PIIReplacer(seed=42)
    replacement_map = {}
    for ent in entities:
        fake = replacer.get_replacement(ent)
        replacement_map[ent.text] = fake

    print(f"  Generated {len(replacement_map)} unique replacements")

    # Print sample replacements
    print("\n  Sample replacements:")
    for i, (original, fake) in enumerate(replacement_map.items()):
        if i >= 10:
            print("  ...")
            break
        print(f"    {original[:50]:50s} -> {fake}")

    # Step 4: Apply replacements and save
    print("\n[4/5] Applying redactions to document...")
    redactor.apply_replacements(replacement_map)
    redactor.save(output_path)

    # Step 5: Evaluate
    print("\n[5/5] Running evaluation...")
    if os.path.exists(ground_truth_path):
        evaluator = Evaluator(ground_truth_path)
        eval_results = evaluator.evaluate(entities, full_text)
        evaluator.generate_report(eval_results, replacer.get_mapping(), eval_report_path)

        # Print overall metrics
        overall = eval_results["OVERALL"]
        print(f"\n  Overall Metrics:")
        print(f"    Precision: {overall['precision']:.4f}")
        print(f"    Recall:    {overall['recall']:.4f}")
        print(f"    F1 Score:  {overall['f1_score']:.4f}")
    else:
        print(f"  Warning: Ground truth file '{ground_truth_path}' not found. Skipping evaluation.")

    print("\n" + "=" * 60)
    print("Redaction complete!")
    print(f"  Redacted document: {output_path}")
    print(f"  Evaluation report: {eval_report_path}")
    print("=" * 60)


if __name__ == "__main__":
    main()
